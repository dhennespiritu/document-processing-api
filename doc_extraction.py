import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from datetime import datetime
import mimetypes

import fitz  # PyMuPDF
from docx import Document
from bs4 import BeautifulSoup
import chardet
import magic
from azure.core.exceptions import AzureError  # Add this import

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ProcessedDocument:
    """Container for processed document data"""
    content: str
    metadata: Dict
    source_file: str
    file_type: str
    page_count: Optional[int] = None
    processing_time: Optional[float] = None
    errors: List[str] = None

class DocumentProcessor:
    """Main document processing class for ingestion and preprocessing"""
    
    def __init__(self, supported_formats: Optional[List[str]] = None):
        self.supported_formats = supported_formats or [
            '.pdf', '.docx', '.doc', '.txt', '.html', '.htm', '.md', '.rtf'
        ]
        self.max_file_size = 50 * 1024 * 1024  # 50MB default limit
        
    def validate_file(self, file_path: Path) -> Tuple[bool, str]:
        """Validate file before processing"""
        print(f"Validating file", {file_path})
        try:
            # Check if file exists
            if not file_path.exists():
                return False, f"File does not exist: {file_path}"
            
            # Check file size
            file_size = file_path.stat().st_size
            if file_size > self.max_file_size:
                return False, f"File too large: {file_size} bytes (max: {self.max_file_size})"
            
            # Check file extension
            if file_path.suffix.lower() not in self.supported_formats:
                return False, f"Unsupported format: {file_path.suffix}"
            
            # Check if file is readable
            if not os.access(file_path, os.R_OK):
                return False, f"File not readable: {file_path}"
            
            return True, "Valid file"
            
        except Exception as e:
            return False, f"Validation error: {str(e)}"
    
    def detect_encoding(self, file_path: Path) -> str:
        """Detect file encoding for text files"""
        print(f"detecting encoding for file", {file_path})
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Sample first 10KB
                result = chardet.detect(raw_data)
                encoding = result.get('encoding', 'utf-8')
                confidence = result.get('confidence', 0)
                
                # Fallback to utf-8 if confidence is too low
                if confidence < 0.7:
                    logger.warning(f"Low confidence ({confidence}) for encoding detection. Using utf-8.")
                    encoding = 'utf-8'
                    
                return encoding
        except Exception as e:
            logger.warning(f"Encoding detection failed: {e}. Using utf-8.")
            return 'utf-8'
    
    def extract_pdf_content(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract content from PDF files with multiple fallback modes"""
        try:
            doc = fitz.open(file_path)
            content = ""
            metadata = {
                'page_count': len(doc),
                **doc.metadata
            }

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Try standard text extraction
                page_text = page.get_text("text") or ""

                # Fallback: blocks
                if not page_text.strip():
                    blocks = page.get_text("blocks") or []
                    block_texts = []
                    for b in blocks:
                        if isinstance(b, (list, tuple)) and len(b) > 4:
                            block_texts.append(b[4])
                    page_text = "\n".join(block_texts)

                # Fallback: dict mode
                if not page_text.strip():
                    text_dict = page.get_text("dict") or {}
                    block_texts = []
                    for block in text_dict.get("blocks", []):
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                block_texts.append(span.get("text", ""))
                    page_text = "\n".join(block_texts)

                if page_text.strip():
                    content += f"\n\n--- Page {page_num + 1} ---\n{page_text}"

            doc.close()
            return content.strip(), metadata

        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")

    
    def extract_docx_content(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract content from DOCX files"""
        print(f"Extracting content from file", {file_path})
        try:
            doc = Document(file_path)
            content = ""
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            # Extract metadata
            props = doc.core_properties
            metadata = {
                'title': props.title or '',
                'author': props.author or '',
                'subject': props.subject or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else '',
                'word_count': len(content.split()) if content else 0
            }
            
            return content.strip(), metadata
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def extract_text_content(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract content from plain text files"""
        print(f"Extracting content from file", {file_path})
        try:
            encoding = self.detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                content = f.read()
            
            metadata = {
                'encoding': encoding,
                'line_count': len(content.splitlines()),
                'word_count': len(content.split()),
                'char_count': len(content)
            }
            
            return content, metadata
            
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def extract_html_content(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract content from HTML files"""
        print(f"Extracting content from file", {file_path})
        try:
            encoding = self.detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Extract text content
            content = soup.get_text()
            
            # Extract metadata from HTML
            title = soup.find('title')
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
            
            metadata = {
                'title': title.get_text() if title else '',
                'description': meta_desc.get('content', '') if meta_desc else '',
                'keywords': meta_keywords.get('content', '') if meta_keywords else '',
                'encoding': encoding
            }
            
            return content, metadata
            
        except Exception as e:
            raise Exception(f"HTML extraction failed: {str(e)}")
    
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        print("Cleaning and normalizing extracted data..")
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove multiple consecutive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove leading/trailing whitespace from lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        # Normalize quotes and dashes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        text = text.replace('—', '-').replace('–', '-')
        
        return text.strip()
    
    def extract_content_by_type(self, file_path: Path) -> Tuple[str, Dict]:
        """Route to appropriate extraction method based on file type"""
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            return self.extract_pdf_content(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_docx_content(file_path)
        elif file_ext in ['.txt', '.md']:
            return self.extract_text_content(file_path)
        elif file_ext in ['.html', '.htm']:
            return self.extract_html_content(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def process_document(self, file_path: str) -> ProcessedDocument:
        """Main method to process a single document"""
        print(f"Processing document: ", {file_path})
        start_time = datetime.now()
        file_path = Path(file_path)
        errors = []
        
        try:
            # Validate file
            is_valid, validation_msg = self.validate_file(file_path)
            if not is_valid:
                raise ValueError(validation_msg)
            
            logger.info(f"Processing: {file_path.name}")
            
            # Extract content based on file type
            raw_content, file_metadata = self.extract_content_by_type(file_path)

            # Clean the extracted text
            cleaned_content = self.clean_text(raw_content)
            
            if not cleaned_content:
                raise ValueError("No content extracted from document")
            
            # Prepare comprehensive metadata
            file_stats = file_path.stat()
            metadata = {
                'filename': file_path.name,
                'file_extension': file_path.suffix.lower(),
                'file_size': file_stats.st_size,
                'processed_at': datetime.now().isoformat(),
                'content_length': len(cleaned_content),
                'word_count': len(cleaned_content.split()),
                **file_metadata  # Merge file-specific metadata
            }
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessedDocument(
                content=cleaned_content,
                metadata=metadata,
                source_file=str(file_path),
                file_type=file_path.suffix.lower(),
                processing_time=processing_time,
                errors=errors if errors else None
            )
            
        except Exception as e:
            error_msg = f"Processing failed for {file_path.name}: {str(e)}"
            logger.error(error_msg)
            errors.append(error_msg)
            
            return ProcessedDocument(
                content="",
                metadata={'filename': file_path.name, 'error': str(e)},
                source_file=str(file_path),
                file_type=file_path.suffix.lower() if file_path.suffix else 'unknown',
                errors=errors
            )
    
    def process_directory(self, directory_path: str, recursive: bool = True) -> List[ProcessedDocument]:
        """Process all supported documents in a directory"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists() or not directory_path.is_dir():
            raise ValueError(f"Invalid directory: {directory_path}")
        
        processed_docs = []
        
        # Get all files
        if recursive:
            pattern = "**/*"
        else:
            pattern = "*"
        
        for file_path in directory_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    processed_doc = self.process_document(str(file_path))
                    processed_docs.append(processed_doc)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    # Continue processing other files
                    continue
        
        logger.info(f"Processed {len(processed_docs)} documents from {directory_path}")
        return processed_docs
    
    def get_processing_stats(self, processed_docs: List[ProcessedDocument]) -> Dict:
        """Generate processing statistics"""
        total_docs = len(processed_docs)
        successful_docs = len([doc for doc in processed_docs if doc.content])
        failed_docs = total_docs - successful_docs
        
        total_content_length = sum(len(doc.content) for doc in processed_docs if doc.content)
        total_processing_time = sum(doc.processing_time for doc in processed_docs if doc.processing_time)
        
        file_types = {}
        for doc in processed_docs:
            file_type = doc.file_type
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            'total_documents': total_docs,
            'successful_extractions': successful_docs,
            'failed_extractions': failed_docs,
            'success_rate': (successful_docs / total_docs * 100) if total_docs > 0 else 0,
            'total_content_length': total_content_length,
            'average_content_length': total_content_length / successful_docs if successful_docs > 0 else 0,
            'total_processing_time': total_processing_time,
            'average_processing_time': total_processing_time / total_docs if total_docs > 0 else 0,
            'file_types_processed': file_types
        }

# Advanced text cleaning utilities
class TextCleaner:
    """Advanced text cleaning and normalization utilities"""
    
    @staticmethod
    def remove_headers_footers(text: str, threshold: int = 3) -> str:
        """Remove repeated headers/footers that appear on multiple pages"""
        lines = text.split('\n')
        line_counts = {}
        
        # Count occurrences of each line
        for line in lines:
            clean_line = line.strip()
            if len(clean_line) > 10:  # Only consider substantial lines
                line_counts[clean_line] = line_counts.get(clean_line, 0) + 1
        
        # Remove lines that appear too frequently (likely headers/footers)
        repeated_lines = {line for line, count in line_counts.items() if count >= threshold}
        
        filtered_lines = []
        for line in lines:
            if line.strip() not in repeated_lines:
                filtered_lines.append(line)
        
        return '\n'.join(filtered_lines)
    
    @staticmethod
    def normalize_spacing(text: str) -> str:
        """Normalize various types of spacing issues"""
        # Fix common spacing issues
        text = re.sub(r'([.!?])\s*\n\s*([A-Z])', r'\1 \2', text)  # Fix sentence breaks
        text = re.sub(r'\s*\n\s*\n\s*', '\n\n', text)  # Normalize paragraph breaks
        text = re.sub(r'[ \t]+', ' ', text)  # Normalize internal spacing
        text = re.sub(r'\n[ \t]+', '\n', text)  # Remove leading spaces on lines
        
        return text
    
    @staticmethod
    def fix_hyphenation(text: str) -> str:
        """Fix line-break hyphenation issues"""
        # Fix hyphenated words split across lines
        text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)
        return text
    
    @staticmethod
    def extract_structure(text: str) -> Dict:
        """Extract document structure information"""
        lines = text.split('\n')
        
        # Detect potential headers (lines that are short and may be titles)
        potential_headers = []
        for i, line in enumerate(lines):
            line = line.strip()
            if line and len(line) < 100 and not line.endswith('.'):
                # Check if next line is longer (body text)
                if i + 1 < len(lines) and len(lines[i + 1].strip()) > len(line):
                    potential_headers.append((i, line))
        
        # Extract bullet points and lists
        bullet_patterns = [r'^\s*[•\-\*]\s+', r'^\s*\d+\.\s+', r'^\s*[a-zA-Z]\.\s+']
        list_items = []
        
        for line in lines:
            for pattern in bullet_patterns:
                if re.match(pattern, line):
                    list_items.append(line.strip())
                    break
        
        return {
            'potential_headers': potential_headers,
            'list_items': list_items,
            'paragraph_count': len([line for line in lines if len(line.strip()) > 50]),
            'average_line_length': sum(len(line) for line in lines) / len(lines) if lines else 0
        }

# Example usage and testing
# def main():
#     """Example usage of the document processor"""
    
#     # Initialize processor
#     processor = DocumentProcessor()
#     cleaner = TextCleaner()
    
#     # Example: Process a single file
#     try:
#         # Replace with your actual file path
#         file_path = Path(r"C:\Users\dv146ms\Downloads\Invoice_000001.pdf").resolve()

        
#         if Path(file_path).exists():
#             processed_doc = processor.process_document(file_path)
#             if processed_doc.content:
#                 print(f"Successfully processed: {processed_doc.source_file}")
#                 print(f"Content length: {len(processed_doc.content)} characters")
#                 print(f"Processing time: {processed_doc.processing_time:.2f} seconds")
#                 print(f"Metadata: {processed_doc.metadata}")
                
#                 # Apply advanced cleaning
#                 enhanced_content = cleaner.remove_headers_footers(processed_doc.content)
#                 enhanced_content = cleaner.normalize_spacing(enhanced_content)
#                 enhanced_content = cleaner.fix_hyphenation(enhanced_content)
                
#                 #structure_info = cleaner.extract_structure(enhanced_content)
#                 #print(f"Document structure: {structure_info}")
#                 print (enhanced_content)
                
#             else:
#                 print(f"Failed to process: {processed_doc.errors}")
        
#         # Example: Process a directory
#         # processed_docs = processor.process_directory("./documents", recursive=True)
#         # stats = processor.get_processing_stats(processed_docs)
#         # print(f"Processing statistics: {stats}")
        
#     except Exception as e:
#         logger.error(f"Processing error: {e}")

# if __name__ == "__main__":
#     main()
